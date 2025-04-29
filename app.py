import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
import base64

# Funzione per generare un nome di file univoco
def generate_unique_filename(nome_cliente):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f'{timestamp}_{nome_cliente}.docx'

# Funzione per aggiungere una linea vuota
def add_empty_line(paragraph, count=1):
    for _ in range(count):
        run = paragraph.add_run()
        run.add_break()

# Funzione per creare il file docx del preventivo
def create_preventivo_docx(nome_cliente, tipo_sito, piattaforma, seo, hosting, altro_sito, descrizione_personalizzata, totale, costo_base, costo_piattaforma, costo_sito, costo_seo, costo_hosting):
    # Creazione documento Word
    filename = generate_unique_filename(nome_cliente)
    document = Document()

    # Stile del preventivo (font e layout)
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    # Intestazione cliente e data
    header = document.sections[0].header
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = "Cliente: " + nome_cliente
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    header_paragraph = header.add_paragraph()
    header_paragraph.text = "Data: " + datetime.now().strftime("%d/%m/%Y")
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Preventivo titolo
    document.add_heading(f'Preventivo per {nome_cliente}', level=1)

    # Tabella preventivo
    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Descrizione'
    hdr_cells[1].text = 'Quantità'
    hdr_cells[2].text = 'Prezzo'
    hdr_cells[3].text = 'Subtotale'

    # Riga 0 (costo base)
    row_cells = table.add_row().cells
    row_cells[0].text = 'Costo base'
    row_cells[1].text = '1'
    row_cells[2].text = f'{costo_base}€'
    row_cells[3].text = f'{costo_base}€'

    # Riga 1
    row_cells = table.add_row().cells
    row_cells[0].text = 'Creazione sito ' + tipo_sito
    row_cells[1].text = '1'
    row_cells[2].text = f'{costo_sito}€'
    row_cells[3].text = f'{costo_sito}€'

    # Riga 2
    row_cells = table.add_row().cells
    row_cells[0].text = 'Piattaforma ' + piattaforma
    row_cells[1].text = '1'
    row_cells[2].text = f'{costo_piattaforma}€'
    row_cells[3].text = f'{costo_piattaforma}€'

    # Riga 3 (SEO)
    row_cells = table.add_row().cells
    row_cells[0].text = 'SEO'
    row_cells[1].text = '1'
    row_cells[2].text = f'{costo_seo}€'
    row_cells[3].text = f'{costo_seo}€'

    # Riga 4 (Hosting)
    row_cells = table.add_row().cells
    row_cells[0].text = 'Hosting'
    row_cells[1].text = '1'
    row_cells[2].text = f'{costo_hosting}€'
    row_cells[3].text = f'{costo_hosting}€'

    # Riga 5 (Descrizione personalizzata)
    if descrizione_personalizzata:
        row_cells = table.add_row().cells
        row_cells[0].text = descrizione_personalizzata
        row_cells[1].text = '1'
        row_cells[2].text = '0€'
        row_cells[3].text = '0€'

    # Aggiunta totale
    paragraph = document.add_paragraph()
    paragraph.add_run(f'Totale: {totale}€').bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Salva il documento
    document.save(filename)
    return filename

# Funzione per consentire il download del file
def get_binary_file_downloader_html(bin_file, file_label='File'):
    with open(bin_file, 'rb') as f:
        data = f.read()
    b64 = base64.b64encode(data).decode()
    href = f'<a href="data:application/octet-stream;base64,{b64}" download="{os.path.basename(bin_file)}">{file_label}</a>'
    return href

# Interfaccia utente Streamlit
st.set_page_config(page_title="Generatore Preventivi", page_icon="💰", layout="wide")

# Titolo dell'app
st.title("Generatore di Preventivi")
st.subheader("Compila il form per generare un preventivo personalizzato")

# Form per la raccolta dati
with st.form("preventivo_form"):
    col1, col2 = st.columns(2)
    
    with col1:
        nome_cliente = st.text_input("Nome Cliente", placeholder="Inserisci il nome del cliente")
        tipo_sito = st.selectbox("Tipo di Sito", options=["blog", "e-commerce", "portfolio", "altro"])
        if tipo_sito == "altro":
            altro_sito = st.text_input("Specifica altro tipo di sito")
        else:
            altro_sito = ""
        
        piattaforma = st.selectbox("Piattaforma", options=["WordPress", "Da zero"])
    
    with col2:
        seo = st.radio("SEO", options=["si", "no"])
        hosting = st.radio("Hosting", options=["si", "no"])
        descrizione_personalizzata = st.text_area("Descrizione Personalizzata", placeholder="Aggiungi dettagli o note specifiche...")
    
    # Bottone per generare il preventivo
    submitted = st.form_submit_button("Genera Preventivo")
    
    if submitted:
        # Calcolo del preventivo
        costo_base = 500
        costo_piattaforma = 300 if piattaforma == "WordPress" else 600
        costo_sito = {
            'blog': 200,
            'e-commerce': 1000,
            'portfolio': 400,
            'altro': 600
        }.get(tipo_sito, 0)
        costo_seo = 200 if seo == "si" else 0
        costo_hosting = 100 if hosting == "si" else 0

        totale = costo_base + costo_piattaforma + costo_sito + costo_seo + costo_hosting
        
        # Genera documento
        filename = create_preventivo_docx(
            nome_cliente, tipo_sito, piattaforma, seo, hosting, 
            altro_sito, descrizione_personalizzata, 
            totale, costo_base, costo_piattaforma, costo_sito, costo_seo, costo_hosting
        )
        
        # Mostra il riepilogo del preventivo
        st.success(f"Preventivo per {nome_cliente} generato con successo!")
        
        # Mostra una card con il riepilogo
        st.subheader("Riepilogo del preventivo")
        
        col1, col2 = st.columns(2)
        with col1:
            st.markdown(f"**Cliente:** {nome_cliente}")
            st.markdown(f"**Tipo di sito:** {tipo_sito}")
            st.markdown(f"**Piattaforma:** {piattaforma}")
        
        with col2:
            st.markdown(f"**SEO:** {seo}")
            st.markdown(f"**Hosting:** {hosting}")
            st.markdown(f"**Totale:** €{totale}")
        
        # Aggiungi dettagli costi
        with st.expander("Dettaglio costi"):
            st.markdown(f"- Costo base: €{costo_base}")
            st.markdown(f"- Costo sito {tipo_sito}: €{costo_sito}")
            st.markdown(f"- Costo piattaforma {piattaforma}: €{costo_piattaforma}")
            st.markdown(f"- Costo SEO: €{costo_seo}")
            st.markdown(f"- Costo hosting: €{costo_hosting}")
        
        # Offri download del file
        st.markdown("### Download")
        st.markdown(get_binary_file_downloader_html(filename, 'Scarica il documento Word'), unsafe_allow_html=True)
        
        # Aggiungi opzione per eliminare il file dopo il download
        if st.button("Elimina file dopo il download"):
            try:
                os.remove(filename)
                st.success("File eliminato con successo!")
            except:
                st.error("Errore nell'eliminazione del file.")

# Footer
st.markdown("---")
st.markdown("© 2025 - Generatore di Preventivi")